"""Gera .docx (Word) do projeto de pesquisa para submissao com a
Declaracao de Responsabilidade (Art. 8 da Res. 200/2021/CONSU Unifesp).

Formato Word editavel — o mestrando pode copiar diretamente e formatar
conforme o padrao do orgao / secretaria.

Uso:
    python _gerar_projeto_docx.py
    -> produz: docs/ativo/projeto_declaracao_responsabilidade.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


NAVY = RGBColor(0x1F, 0x2A, 0x4E)
GRAY_DK = RGBColor(0x3D, 0x42, 0x4E)
GRAY_MD = RGBColor(0x70, 0x76, 0x82)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_header_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_MD


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)


def _add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.italic = True
    run.font.color.rgb = GRAY_DK
    p.paragraph_format.space_after = Pt(18)


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_body(doc: Document, segments) -> None:
    """Adiciona parágrafo com segmentos: string ou (texto, estilo).

    Estilos: 'b' (bold), 'i' (italic).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35

    if isinstance(segments, str):
        segments = [segments]

    for seg in segments:
        if isinstance(seg, tuple):
            text, style = seg
        else:
            text, style = seg, None
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = GRAY_DK
        if style == "b":
            run.bold = True
        elif style == "i":
            run.italic = True


def _add_bullet(doc: Document, segments) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35

    if isinstance(segments, str):
        segments = [segments]

    for seg in segments:
        if isinstance(seg, tuple):
            text, style = seg
        else:
            text, style = seg, None
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = GRAY_DK
        if style == "b":
            run.bold = True
        elif style == "i":
            run.italic = True


def _add_kv_table(doc: Document, rows) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(12.0)

    for i, (k, v) in enumerate(rows):
        cell_k = table.rows[i].cells[0]
        cell_v = table.rows[i].cells[1]
        cell_k.width = Cm(5.0)
        cell_v.width = Cm(12.0)

        for cell in (cell_k, cell_v):
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

        # Chave em bold navy
        rk = cell_k.paragraphs[0].add_run(k)
        rk.font.name = "Calibri"
        rk.font.size = Pt(10)
        rk.bold = True
        rk.font.color.rgb = NAVY

        # Valor em cinza
        rv = cell_v.paragraphs[0].add_run(v)
        rv.font.name = "Calibri"
        rv.font.size = Pt(10)
        rv.font.color.rgb = GRAY_DK


def _add_data_table(doc: Document, headers, rows) -> None:
    n_cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        cell = hdr_cells[j]
        _set_cell_bg(cell, "1F2A4E")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Linhas
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRAY_DK


def _add_signature(doc: Document, name: str, role: str) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 55)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(role)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.italic = True


def build_docx(out: Path) -> None:
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cabeçalho institucional
    _add_header_paragraph(doc, "UNIVERSIDADE FEDERAL DE SÃO PAULO — UNIFESP")
    _add_header_paragraph(doc, "Instituto de Ciência e Tecnologia — ICT")
    _add_header_paragraph(doc, "Programa de Pós-Graduação em Ciência da Computação")

    doc.add_paragraph()

    # Título
    _add_title(doc, "Projeto de Pesquisa — Anexo à Declaração de Responsabilidade")
    _add_subtitle(
        doc,
        "Documento formal do projeto para submissão junto à Declaração de "
        "Responsabilidade prevista no parágrafo único do Art. 8º da Resolução "
        "nº 200/2021/CONSELHO UNIVERSITÁRIO da Universidade Federal de São Paulo.",
    )

    # 1. Identificação
    _add_h1(doc, "1. Identificação")
    _add_kv_table(
        doc,
        [
            ("Título do projeto", "Mitigação de Viés Racial em Classificação Facial via Condicionamento por Tom de Pele em Arquiteturas Profundas"),
            ("Mestrando", "Marcello Ozzetti"),
            ("Orientador / Pesquisador Responsável", "Prof. Dr. Marcos Quiles"),
            ("Programa", "Mestrado em Ciência da Computação"),
            ("Instituição", "Universidade Federal de São Paulo — Instituto de Ciência e Tecnologia (Unifesp / ICT)"),
            ("Linha de pesquisa", "Visão Computacional e Equidade Algorítmica"),
            ("Período previsto", "julho de 2026 a outubro de 2026 (qualificação)"),
        ],
    )

    # 2. Resumo
    _add_h1(doc, "2. Resumo")
    _add_body(
        doc,
        [
            "Sistemas modernos de reconhecimento facial apresentam disparidade demográfica significativa entre grupos raciais, mesmo quando treinados sobre datasets explicitamente balanceados. No estado da arte atual, o modelo FaceScanPaliGemma sobre o dataset público FairFace atinge F1 de 90 % para faces classificadas como Black, mas apenas 60 % para faces classificadas como Latinx — uma diferença de trinta pontos percentuais que persiste através de múltiplos métodos e arquiteturas. Esta dissertação propõe um pipeline de classificação racial que incorpora tom de pele (escala Monk Skin Tone, dez classes) como sinal auxiliar condicionante via mecanismo arquitetural FiLM (Feature-wise Linear Modulation), e avalia sua capacidade de reduzir essa disparidade sem sacrificar acurácia agregada. ",
            ("A pesquisa é inteiramente computacional, conduzida sobre datasets públicos secundários, sem coleta primária, intervenção ou qualquer forma de envolvimento direto ou indireto de seres humanos.", "b"),
        ],
    )

    # 3. Justificativa
    _add_h1(doc, "3. Justificativa e problema de pesquisa")
    _add_body(
        doc,
        "O reconhecimento facial deixou de ser tecnologia laboratorial e passou a integrar a infraestrutura social — desbloqueio de dispositivos, autenticação bancária, identidade digital, controle de fronteiras e identificação policial. Em escala industrial, o National Institute of Standards and Technology auditou em 2019 cento e oitenta e nove algoritmos de reconhecimento facial sobre dezoito milhões de imagens, documentando diferenciais de dez a cem vezes na taxa de falso positivo entre grupos raciais.",
    )
    _add_body(
        doc,
        "Apesar de mais de uma década de pesquisa em equidade algorítmica, a disparidade persiste. O baseline ResNet-34 sobre FairFace, o FaceScanPaliGemma via VLM, e mesmo modelos com balanceamento explícito de dados continuam apresentando F1 aproximadamente 60 % para a categoria Latinx / Hispanic. Trabalhos recentes sugerem que essa disparidade não é apenas artefato algorítmico, mas reflete heterogeneidade fenotípica intra-categorial que a rotulagem racial monolítica do FairFace não captura adequadamente.",
    )
    _add_body(
        doc,
        [
            "O problema de pesquisa desta dissertação é, portanto: ",
            ("condicionar arquiteturalmente um classificador racial pelo tom de pele reduz a disparidade demográfica sem sacrificar acurácia agregada, e permite decompor o erro em componentes fenotípico e algorítmico?", "b"),
        ],
    )

    # 4. Objetivos
    _add_h1(doc, "4. Objetivos")
    _add_h2(doc, "4.1 Objetivo geral")
    _add_body(
        doc,
        "Desenvolver e avaliar um pipeline de classificação racial em imagens faciais que incorpora tom de pele (escala Monk Skin Tone) como sinal auxiliar condicionante via mecanismo arquitetural FiLM, com o propósito de mitigar viés racial demonstrável no estado da arte atual — particularmente a disparidade severa entre classes raciais bem representadas e classes sub-representadas documentada em modelos como FaceScanPaliGemma sobre o dataset FairFace.",
    )

    _add_h2(doc, "4.2 Objetivos específicos")
    oes = [
        [("1. Auditoria fenotípica do FairFace", "b"), " — quantificar a distribuição cruzada MST × classes raciais sobre o FairFace via SkinToneNet pré-treinado, entregando a primeira matriz pública dessa distribuição."],
        [("2. Avaliação metodológica de classificadores MST", "b"), " — comparar SkinToneNet, Casual Conversations baseline e alternativas disponíveis, com ", ("sensitivity analysis", "i"), " a dois ou três classificadores alternativos para mitigar risco de propagação de viés."],
        [("3. Race classifier com conditioning arquitetural", "b"), " — treinar ConvNeXt-T ", ("fine-tuned", "i"), " em FairFace, com camadas FiLM por estágio recebendo o vetor MST como contexto."],
        [("4. Comparação sistemática contra baselines de mitigação", "b"), " — avaliar o pipeline contra seis baselines (ResNet-34, ConvNeXt-T puro, FSCL+, Group DRO, FineFACE, Adversarial debiasing) via triangulação de métricas (DR + ", ("worst-class", "i"), " F1 + EO_h/EOD)."],
        [("5. Fair transferência para face recognition downstream", "b"), " — aplicar o backbone fair em ", ("face recognition", "i"), " em RFW ou BFW, com controle explícito de ", ("pixel information", "i"), " como ", ("confounder", "i"), "."],
        [("6. Síntese decompositiva", "b"), " — combinar resultados dos objetivos 1 e 5 para quantificar componente fenotípico (irredutível) versus componente algorítmico (mitigável) do erro Latinx."],
    ]
    for oe in oes:
        _add_bullet(doc, oe)

    doc.add_page_break()

    # 5. Metodologia
    _add_h1(doc, "5. Metodologia")
    _add_h2(doc, "5.1 Natureza da pesquisa")
    _add_body(
        doc,
        [
            "A pesquisa é ",
            ("puramente computacional", "b"),
            ". Consiste em treinamento e avaliação de modelos de aprendizado profundo sobre datasets de imagens faciais já existentes e publicamente disponíveis; análise estatística de métricas de desempenho estratificadas por grupo demográfico; comparação sistemática entre configurações arquiteturais; e síntese analítica dos resultados.",
        ],
    )

    _add_h2(doc, "5.2 Datasets utilizados — todos secundários e públicos")
    _add_data_table(
        doc,
        ["Dataset", "Origem", "Licença", "Uso na pesquisa"],
        [
            ["FairFace", "Kärkkäinen & Joo (2021), UCLA", "CC BY 4.0", "Dataset principal — 108.501 imagens em 7 categorias raciais"],
            ["RFW", "Wang et al. (2019), BUPT", "Research-only", "Fair transferência downstream"],
            ["BFW", "Robinson et al. (2020), RIT", "CC BY 4.0", "Avaliação alternativa de fair transferência"],
            ["BUPT-Balancedface", "Wang, Zhang, Deng (2022)", "Research-only", "Sensitivity analysis do backbone"],
            ["STW", "Pereira et al. (2026)", "Research-only", "Base de treino do SkinToneNet (modelo pré-treinado)"],
        ],
    )
    _add_body(doc, "")
    _add_body(
        doc,
        [
            ("Todas as imagens desses datasets foram coletadas, anotadas e disponibilizadas por seus autores originais anteriormente ao início desta pesquisa. Este projeto não coleta, não anota e não distribui qualquer imagem original.", "b"),
        ],
    )

    _add_h2(doc, "5.3 Pipeline metodológico em seis etapas")
    etapas = [
        [("Etapa 1 — Classificador MST", "b"), ": uso do SkinToneNet pré-treinado como insumo, mediante validação de concordância interna em subset estratificado."],
        [("Etapa 2 — Auditoria FairFace", "b"), ": aplicação do SkinToneNet sobre o FairFace validation set e construção da matriz pública MST × raça."],
        [("Etapa 3 — Race classifier com conditioning", "b"), ": treinamento do ConvNeXt-T em FairFace train, com camadas FiLM recebendo o vetor MST como contexto arquitetural."],
        [("Etapa 4 — Comparação contra baselines", "b"), ": avaliação do pipeline contra seis baselines de mitigação, com triangulação métrica."],
        [("Etapa 5 — Fair transferência", "b"), ": aplicação do backbone fair em face recognition downstream (RFW ou BFW)."],
        [("Etapa 6 — Síntese decompositiva", "b"), ": quantificação do componente fenotípico versus algorítmico do erro Latinx."],
    ]
    for e in etapas:
        _add_bullet(doc, e)

    _add_h2(doc, "5.4 Rigor experimental")
    _add_body(
        doc,
        [
            "Todos os experimentos serão conduzidos com múltiplas sementes aleatórias (mínimo três: 42, 1, 2), comparação pareada entre configurações, intervalos de confiança via ",
            ("bootstrap", "i"),
            " e reporte estratificado por raça e por interseção raça × gênero. As métricas seguirão triangulação (DR + ",
            ("worst-class", "i"),
            " F1 + EO_h/EOD) para endereçar o Teorema da Impossibilidade de Kleinberg.",
        ],
    )

    _add_h2(doc, "5.5 Validação humana — modalidade interna")
    _add_body(
        doc,
        [
            "Para o Objetivo Específico 1, será conduzida validação de concordância entre o SkinToneNet e anotação manual em subset estratificado de aproximadamente duzentas a trezentas imagens do FairFace. ",
            ("Esta validação será realizada exclusivamente pela equipe acadêmica do projeto", "b"),
            " — Mestrando e Orientador (e Coorientador, se designado). ",
            ("Não haverá contratação de anotadores externos, não haverá uso de plataformas de crowdsourcing, e não haverá coleta de qualquer dado pessoal de indivíduos externos à equipe.", "b"),
        ],
    )

    doc.add_page_break()

    # 6. Aspectos éticos
    _add_h1(doc, "6. Aspectos éticos — enquadramento no Art. 8º da Resolução 200/2021")
    _add_body(
        doc,
        [
            "Este projeto ",
            ("não envolve, direta ou indiretamente, seres humanos como sujeitos de pesquisa", "b"),
            ", conforme a definição estabelecida pela Resolução nº 200/2021 do Conselho Universitário da Universidade Federal de São Paulo e pelas Resoluções nº 466/2012 e nº 510/2016 do Conselho Nacional de Saúde. Justifica-se:",
        ],
    )

    _add_h2(doc, "6.1 Ausência de coleta primária")
    _add_body(
        doc,
        [
            "O projeto ",
            ("não coleta", "b"),
            ", em nenhuma etapa, imagens de rostos, dados biométricos, dados demográficos, opiniões, comportamento ou qualquer outra informação de seres humanos identificáveis ou identificados. Todas as imagens utilizadas são secundárias, provenientes de datasets construídos e disponibilizados publicamente por terceiros anteriormente ao início desta pesquisa.",
        ],
    )

    _add_h2(doc, "6.2 Ausência de intervenção")
    _add_body(
        doc,
        [
            "O projeto ",
            ("não realiza intervenção", "b"),
            " de qualquer natureza sobre seres humanos — não há entrevistas, questionários, experimentos, testes psicológicos, exames físicos, uso de dispositivos, alteração de ambiente ou qualquer procedimento que produza efeito sobre pessoas físicas.",
        ],
    )

    _add_h2(doc, "6.3 Ausência de envolvimento indireto via crowdsourcing")
    _add_body(
        doc,
        [
            "Diferentemente de projetos que empregam plataformas de crowdsourcing (Amazon Mechanical Turk, Prolific, Toloka, entre outras) para anotação de dados, este projeto ",
            ("não contrata anotadores externos", "b"),
            ". A validação manual prevista no Objetivo Específico 1 será conduzida exclusivamente por membros da equipe acadêmica do projeto, enquadrando-se como atividade de pesquisa científica interna e não como recrutamento de participantes.",
        ],
    )

    _add_h2(doc, "6.4 Ausência de identificação de sujeitos")
    _add_body(
        doc,
        [
            "As análises produzidas nesta pesquisa referem-se a ",
            ("grupos demográficos agregados", "b"),
            " (categorias raciais, tons de pele) e não a indivíduos identificáveis. Nenhum resultado publicado permitirá identificar pessoas específicas presentes nos datasets utilizados.",
        ],
    )

    _add_h2(doc, "6.5 Ausência de uso de animais vertebrados vivos")
    _add_body(
        doc,
        [
            "O projeto ",
            ("não utiliza animais vertebrados vivos", "b"),
            " de qualquer natureza, não se enquadrando na competência da Comissão de Ética no Uso de Animais (CEUA).",
        ],
    )

    _add_h2(doc, "6.6 Enquadramento formal")
    _add_body(
        doc,
        [
            "Pelo exposto, este projeto se enquadra no ",
            ("Art. 8º da Resolução nº 200/2021/CONSU", "b"),
            " — \"Os projetos de pesquisa que não envolvem seres humanos, direta e indiretamente, nem animais vertebrados vivos, estão dispensados de cadastro\" — e é objeto da ",
            ("Declaração de Responsabilidade", "b"),
            " prevista no parágrafo único do referido artigo, assinada pelo estudante, pelo orientador e pelo chefe do Departamento ao qual o orientador está vinculado.",
        ],
    )

    # 7. Cronograma
    _add_h1(doc, "7. Cronograma")
    _add_data_table(
        doc,
        ["Marco", "Data", "Descrição"],
        [
            ["Primeira revisão da qualificação ao orientador", "15/jul/2026", "Documento completo em LaTeX / Overleaf"],
            ["Pedido formal de qualificação ao Programa", "30/jul/2026", "Prazo regimental do PPG-CC / ICT"],
            ["Defesa da qualificação", "outubro/2026", "Sujeita à prorrogação já solicitada de dois meses"],
            ["Experimentos e escrita da dissertação", "nov/2026 – mai/2027", "—"],
            ["Defesa da dissertação", "2º semestre/2027", "—"],
        ],
    )

    doc.add_page_break()

    # 8. Referências
    _add_h1(doc, "8. Referências principais")

    _add_h2(doc, "Datasets e benchmarks")
    for r in [
        "Kärkkäinen & Joo (2021). FairFace: Face Attribute Dataset. WACV.",
        "Wang et al. (2019). Racial Faces in the Wild — RFW. ICCV.",
        "Robinson et al. (2020). Face Recognition: Too Bias, or Not Too Bias? CVPRW.",
        "Hazirbas et al. (2021). Casual Conversations. CVPRW.",
    ]:
        _add_bullet(doc, r)

    _add_h2(doc, "Fundamentação teórica em fairness")
    for r in [
        "Buolamwini & Gebru (2018). Gender Shades. FAT*.",
        "Hardt et al. (2016). Equality of Opportunity in Supervised Learning. NeurIPS.",
        "Kleinberg et al. (2017). Inherent Trade-Offs in the Fair Determination of Risk Scores. ITCS.",
        "Madras et al. (2018). Learning Adversarially Fair and Transferable Representations. ICML.",
    ]:
        _add_bullet(doc, r)

    _add_h2(doc, "Skin tone e conditioning arquitetural")
    for r in [
        "Schumann et al. (2023). Consensus and Subjectivity of Skin Tone Annotation — MST-E. CVPR.",
        "Pereira et al. (2026). Large-Scale Dataset and Benchmark for Skin Tone Classification in the Wild — SkinToneNet. arXiv:2603.02475.",
        "Perez et al. (2018). FiLM: Visual Reasoning with a General Conditioning Layer. AAAI.",
    ]:
        _add_bullet(doc, r)

    _add_h2(doc, "Refutação e diálogo crítico")
    for r in [
        "Pangelinan et al. (2023). Analyzing Bias in Race Classification. FAccT.",
        "AlDahoul et al. (2024). FaceScanPaliGemma. arXiv.",
    ]:
        _add_bullet(doc, r)

    _add_h2(doc, "Diversidade fenotípica intra-Latinx")
    for r in [
        "Telles (2014). Pigmentocracies: Ethnicity, Race, and Color in Latin America. UNC Press.",
        "Bryc et al. (2015). The Genetic Ancestry of African Americans, Latinos, and European Americans across the United States. AJHG.",
        "Pew Research (2017). Hispanic Identity Fades Across Generations.",
    ]:
        _add_bullet(doc, r)

    _add_body(
        doc,
        [
            ("Bibliografia consolidada de 104 referências disponível no repositório do projeto (arquivo referencias.bib).", "i"),
        ],
    )

    # 9. Local
    _add_h1(doc, "9. Local de execução")
    _add_body(
        doc,
        [
            "Toda a pesquisa será conduzida em ",
            ("ambiente computacional próprio", "b"),
            " do mestrando e em infraestrutura de computação disponibilizada pela Unifesp / ICT ou por serviços de computação em nuvem contratados individualmente (Google Colab Pro, AWS educacional ou equivalente). ",
            ("Não há uso de laboratórios institucionais que envolvam seres humanos ou animais.", "b"),
        ],
    )

    # 10. Declaração final
    _add_h1(doc, "10. Declaração final")
    _add_body(
        doc,
        "Declaro, para os devidos fins, que o projeto de pesquisa aqui descrito não envolve, direta ou indiretamente, seres humanos como sujeitos de pesquisa, nem faz uso de animais vertebrados vivos, enquadrando-se no Art. 8º da Resolução nº 200/2021 do Conselho Universitário da Universidade Federal de São Paulo. Comprometo-me a submeter novo protocolo ao Comitê de Ética em Pesquisa da Unifesp caso, no curso da pesquisa, sobrevenha qualquer alteração de escopo que envolva seres humanos ou animais.",
    )

    # Assinaturas
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("São Paulo, 8 de julho de 2026.")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)

    _add_signature(doc, "Marcello Ozzetti", "Mestrando")
    _add_signature(doc, "Prof. Dr. Marcos Quiles", "Orientador / Pesquisador Responsável")

    doc.save(str(out))


def main() -> None:
    here = Path(__file__).parent
    out = here / "projeto_declaracao_responsabilidade.docx"
    build_docx(out)
    print(f"DOCX gerado: {out}")
    print(f"Tamanho: {out.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
